"""
report_generator.py
---------------------
Builds a polished, downloadable Word (.docx) report summarizing the GST
analysis: KPIs, year-wise table, and the AI-generated (or fallback)
narrative insight text. Returns an in-memory BytesIO buffer so Streamlit
can offer it directly as a download with no temp files.
"""

from __future__ import annotations

import io
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from data_loader import format_indian_rupees

NAVY = RGBColor(0x0A, 0x1F, 0x44)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
GREY = RGBColor(0x55, 0x55, 0x55)


def _set_cell_shading(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _style_heading(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _add_title_page(doc: Document, trade_name: str, stats: dict, gstin):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GST Insight AI")
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("GST Performance & Compliance Report")
    run2.font.size = Pt(16)
    run2.font.color.rgb = GOLD

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f"{trade_name}")
    info_run.font.size = Pt(14)
    info_run.font.bold = True
    info_run.font.color.rgb = NAVY

    if gstin and str(gstin).strip() and str(gstin).lower() != "nan":
        gst_p = doc.add_paragraph()
        gst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gst_run = gst_p.add_run(f"GSTIN: {gstin}")
        gst_run.font.size = Pt(11)
        gst_run.font.color.rgb = GREY

    date_start, date_end = stats.get("date_range", (None, None))
    period_text = ""
    if date_start is not None and date_end is not None:
        period_text = f"Period covered: {date_start.strftime('%b %Y')} to {date_end.strftime('%b %Y')}"
    period_p = doc.add_paragraph()
    period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_run = period_p.add_run(period_text)
    period_run.font.size = Pt(11)
    period_run.font.color.rgb = GREY

    gen_p = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_run = gen_p.add_run(f"Generated on {datetime.now().strftime('%d %B %Y')}")
    gen_run.font.size = Pt(10)
    gen_run.font.italic = True
    gen_run.font.color.rgb = GREY

    doc.add_page_break()


def _add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A227")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    return p


def _add_kpi_table(doc: Document, stats: dict, show_late_fees: bool = False):
    rows = [
        ("Total Sales", format_indian_rupees(stats.get('total_sales', 0), prefix="₹ ")),
        ("Total GST Liability", format_indian_rupees(stats.get('total_gst_liability', 0), prefix="₹ ")),
        ("Average Monthly Sales", format_indian_rupees(stats.get('avg_monthly_sales', 0), prefix="₹ ")),
        ("Total ITC Available", format_indian_rupees(stats.get('total_itc_available', 0), prefix="₹ ")),
        ("Net ITC", format_indian_rupees(stats.get('net_itc', 0), prefix="₹ ")),
        ("ITC Efficiency", f"{stats.get('itc_efficiency', 0):.2f}%"),
        ("Cash Dependency", f"{stats.get('cash_dependency', 0):.2f}%"),
        ("Reverse Charge %", f"{stats.get('rcm_ratio_pct', 0):.2f}%"),
    ]
    if show_late_fees:
        rows.append(("Total Late Fees / Penalty", format_indian_rupees(stats.get('total_late_fee', 0), prefix="₹ ")))

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(3.2), Inches(3.2)]
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for cell, w in zip(hdr, widths):
        cell.width = w
        _set_cell_shading(cell, "0A1F44")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, (label, value) in enumerate(rows):
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value
        for cell, w in zip(row_cells, widths):
            cell.width = w
            if i % 2 == 0:
                _set_cell_shading(cell, "F2F2F2")
    doc.add_paragraph()


def _add_yearly_table(doc: Document, yearly_df: pd.DataFrame):
    cols = ["financial_year", "months_filed", "total_turnover", "total_outward_tax",
            "total_itc_utilized", "total_cash_paid"]
    headers = [
        "Financial Year",
        "Months Filed",
        "Turnover (₹)",
        "Output Tax (₹)\n(Output Tax ÷ Turnover × 100)",
        "ITC Utilized (₹)\n(ITC Utilized ÷ Output Tax × 100)",
        "Cash Paid (₹)"
    ]

    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _set_cell_shading(hdr[i], "0A1F44")
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, (_, row) in enumerate(yearly_df.iterrows()):
        outward_tax_pct = (row['total_outward_tax'] / row['total_turnover'] * 100) if row['total_turnover'] > 0 else 0.0
        itc_utilized_pct = (row['total_itc_utilized'] / row['total_outward_tax'] * 100) if row['total_outward_tax'] > 0 else 0.0

        cells = table.add_row().cells
        values = [
            str(row["financial_year"]),
            str(int(row["months_filed"])),
            format_indian_rupees(row['total_turnover'], prefix=""),
            f"{format_indian_rupees(row['total_outward_tax'], prefix='')}\n({outward_tax_pct:.2f}%)",
            f"{format_indian_rupees(row['total_itc_utilized'], prefix='')}\n({itc_utilized_pct:.2f}%)",
            format_indian_rupees(row['total_cash_paid'], prefix=""),
        ]
        for j, val in enumerate(values):
            cells[j].text = val
            for p in cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
            if i % 2 == 0:
                _set_cell_shading(cells[j], "F2F2F2")
    doc.add_paragraph()


def _add_compliance_table(doc: Document, df: pd.DataFrame):
    if "filing_date" in df.columns and df["filing_date"].notna().any():
        cols = ["month_label", "due_date", "filing_date", "filing_delay_days", "late_fee"]
        headers = ["Return Period", "Due Date", "Filing Date", "Delay", "Penalty (Late Fee)"]

        table = doc.add_table(rows=1, cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            _set_cell_shading(hdr[i], "0A1F44")
            for p in hdr[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for i, (_, row) in enumerate(df.iterrows()):
            cells = table.add_row().cells
            due_dt = row['due_date'].strftime('%d %b %Y') if pd.notna(row['due_date']) else "N/A"
            fil_dt = row['filing_date'].strftime('%d %b %Y') if pd.notna(row['filing_date']) else "N/A"
            delay = int(row['filing_delay_days']) if pd.notna(row['filing_delay_days']) else 0
            delay_str = f"{delay} days" if delay > 0 else "On Time"
            penalty_str = format_indian_rupees(row['late_fee'], prefix="₹ ") if pd.notna(row['late_fee']) else "₹ 0.00"

            values = [
                str(row["month_label"]),
                due_dt,
                fil_dt,
                delay_str,
                penalty_str,
            ]
            for j, val in enumerate(values):
                cells[j].text = val
                for p in cells[j].paragraphs:
                    if j == 4:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.size = Pt(9)
                if i % 2 == 0:
                    _set_cell_shading(cells[j], "F2F2F2")
        
        # Add Total Row
        total_months = len(df)
        total_penalty = df["late_fee"].sum() if "late_fee" in df.columns else 0.0
        
        cells = table.add_row().cells
        values = [
            f"Total ({total_months} Months)",
            "",
            "",
            "",
            format_indian_rupees(total_penalty, prefix="₹ "),
        ]
        for j, val in enumerate(values):
            cells[j].text = val
            for p in cells[j].paragraphs:
                if j == 4:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.bold = True
                    
        doc.add_paragraph()


def _add_narrative(doc: Document, text: str):
    sections_list = [
        "Monthly Sales Trend",
        "GST Liability Trend",
        "ITC Efficiency",
        "Cash Flow Analysis",
        "Reverse Charge Analysis",
        "Filing Compliance",
    ]
    
    paragraphs = text.split("\n\n")
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        matched_section = None
        for sec in sections_list:
            if para.startswith(sec) or para.lower().startswith(sec.lower()):
                matched_section = sec
                break
        
        if matched_section:
            lines = para.split("\n")
            sec_heading = lines[0].strip()
            
            p_sec = doc.add_paragraph()
            r_sec = p_sec.add_run(sec_heading)
            r_sec.font.size = Pt(13)
            r_sec.font.bold = True
            r_sec.font.color.rgb = GOLD
            p_sec.paragraph_format.space_before = Pt(12)
            p_sec.paragraph_format.space_after = Pt(4)
            
            if len(lines) > 1:
                content = "\n".join(lines[1:]).strip()
                if content:
                    p_content = doc.add_paragraph()
                    run_content = p_content.add_run(content)
                    run_content.font.size = Pt(11)
                    p_content.paragraph_format.space_after = Pt(10)
                    p_content.paragraph_format.line_spacing = 1.15
        else:
            p = doc.add_paragraph()
            run = p.add_run(para)
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.15


def _add_chart_image(doc: Document, image_bytes: bytes, caption: str):
    doc.add_picture(io.BytesIO(image_bytes), width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GREY
    doc.add_paragraph()


def build_word_report(
    trade_name,
    gstin,
    stats: dict,
    yearly_df: pd.DataFrame,
    insight_text: str,
    insight_source: str,
    chart_images=None,
    df: pd.DataFrame | None = None,
) -> io.BytesIO:
    """
    chart_images: optional dict of {caption: png_bytes} to embed (e.g.
    exported from Plotly figures via fig.to_image(format='png')).
    """
    doc = Document()
    _style_heading(doc)

    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    _add_title_page(doc, trade_name, stats, gstin)

    _add_section_heading(doc, "Executive Summary")
    show_late_fees = df is not None and "filing_date" in df.columns and df["filing_date"].notna().any()
    _add_kpi_table(doc, stats, show_late_fees=show_late_fees)

    _add_section_heading(doc, "Year-wise Performance")
    _add_yearly_table(doc, yearly_df)

    if df is not None:
        _add_section_heading(doc, "GSTR-3B Filing Compliance & Penalty Details")
        _add_compliance_table(doc, df)

    if chart_images:
        _add_section_heading(doc, "Key Charts")
        for caption, img_bytes in chart_images.items():
            _add_chart_image(doc, img_bytes, caption)

    _add_section_heading(doc, "AI-Generated Narrative Insights")
    source_note = doc.add_paragraph()
    note_run = source_note.add_run(
        "Generated by Sarvam AI" if insight_source == "sarvam_ai"
        else "Generated locally (Sarvam AI unavailable for this session)"
    )
    note_run.font.size = Pt(9)
    note_run.font.italic = True
    note_run.font.color.rgb = GREY
    _add_narrative(doc, insight_text)

    _add_section_heading(doc, "Conclusion")
    conclusion = doc.add_paragraph()
    conclusion_run = conclusion.add_run(
        f"This report summarizes {stats.get('n_periods')} GST return periods for {trade_name}. "
        "All figures are derived directly from the uploaded return data. This report is intended "
        "for internal review and does not constitute tax or legal advice; please consult a qualified "
        "Chartered Accountant before acting on any observation in this report."
    )
    conclusion_run.font.size = Pt(10)
    conclusion_run.font.color.rgb = GREY

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
